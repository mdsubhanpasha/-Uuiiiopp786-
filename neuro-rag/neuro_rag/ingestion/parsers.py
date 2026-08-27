"""
PASHA-NEURO-RAG Document Parsers
Author: Mohammad Subhan Pasha
Supports PDF, DOCX, URL, and Notion documents with fallback mechanisms.
"""

import os
import re
import urllib.request
from typing import List, Dict, Any, Optional
from bs4 import BeautifulSoup

from neuro_rag.ingestion.schemas import Document, DocumentMetadata


class BaseParser:
    def parse(self, source: str, extra_meta: Optional[Dict[str, Any]] = None) -> Document:
        raise NotImplementedError


class PDFParser(BaseParser):
    def parse(self, source: str, extra_meta: Optional[Dict[str, Any]] = None) -> Document:
        file_path = source
        source_name = os.path.basename(file_path)
        content_text = ""

        # Try pypdf first
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"--- Page {i+1} ---\n{text}")
            content_text = "\n\n".join(pages)
        except Exception:
            # Fallback to PyMuPDF / fitz if pypdf fails or uninstalled
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                pages = [f"--- Page {i+1} ---\n{page.get_text()}" for i, page in enumerate(doc)]
                content_text = "\n\n".join(pages)
            except Exception as e:
                raise RuntimeError(f"Failed to parse PDF file {file_path}: {e}")

        meta = DocumentMetadata(
            source_type="pdf",
            source_name=source_name,
            uri=file_path,
            extra=extra_meta or {}
        )
        return Document(content=content_text, metadata=meta)


class DOCXParser(BaseParser):
    def parse(self, source: str, extra_meta: Optional[Dict[str, Any]] = None) -> Document:
        file_path = source
        source_name = os.path.basename(file_path)
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content_text = "\n\n".join(paragraphs)
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX file {file_path}: {e}")

        meta = DocumentMetadata(
            source_type="docx",
            source_name=source_name,
            uri=file_path,
            extra=extra_meta or {}
        )
        return Document(content=content_text, metadata=meta)


class URLParser(BaseParser):
    def parse(self, source: str, extra_meta: Optional[Dict[str, Any]] = None) -> Document:
        url = source
        try:
            req = urllib.request.Request(
                url,
                headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) PashaNeuroRAG/1.0"}
            )
            with urllib.request.urlopen(req, timeout=15) as response:
                html = response.read().decode("utf-8", errors="ignore")

            soup = BeautifulSoup(html, "html.parser")
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.extract()
            text = soup.get_text(separator="\n")
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            clean_text = "\n".join(chunk for chunk in chunks if chunk)
            title = soup.title.string if soup.title else url
        except Exception as e:
            raise RuntimeError(f"Failed to fetch and parse URL {url}: {e}")

        meta = DocumentMetadata(
            source_type="url",
            source_name=str(title).strip(),
            uri=url,
            extra=extra_meta or {}
        )
        return Document(content=clean_text, metadata=meta)


class NotionParser(BaseParser):
    def parse(self, source: str, extra_meta: Optional[Dict[str, Any]] = None) -> Document:
        """
        Notion parser accepts either raw notion content or notion export file path / URL.
        """
        if os.path.isfile(source):
            if source.endswith(".docx"):
                return DOCXParser().parse(source, extra_meta)
            else:
                with open(source, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                source_name = os.path.basename(source)
        else:
            content = source
            source_name = "Notion Page"

        meta = DocumentMetadata(
            source_type="notion",
            source_name=source_name,
            uri=source if os.path.isfile(source) else None,
            extra=extra_meta or {}
        )
        return Document(content=content, metadata=meta)


class DocumentParserFactory:
    _parsers = {
        "pdf": PDFParser(),
        "docx": DOCXParser(),
        "url": URLParser(),
        "notion": NotionParser(),
    }

    @classmethod
    def parse(cls, source_type: str, source: str, extra_meta: Optional[Dict[str, Any]] = None) -> Document:
        st = source_type.lower()
        if st not in cls._parsers:
            raise ValueError(f"Unsupported source type: {source_type}. Supported: {list(cls._parsers.keys())}")
        return cls._parsers[st].parse(source, extra_meta)
